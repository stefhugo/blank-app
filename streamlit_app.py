import os
import PyPDF2
from docx import Document
from io import BytesIO
import logging
import streamlit as st
import google.generativeai as genai
import json
from pydantic import BaseModel, Field
from typing import List, Optional, Dict

# Set page config to use wide layout
st.set_page_config(layout="wide")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Setup upload directory
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
logging.info(f"Ensuring upload directory exists at: {UPLOAD_FOLDER}")

# Define Pydantic models for the JSON schema
class PersonalInformation(BaseModel):
    full_name: Optional[str] = Field(None, description="Full name of the candidate")
    preferred_name: Optional[str] = Field(None, description="Preferred name of the candidate")
    citizenship_status: Optional[str] = Field(None, description="Citizenship status of the candidate")
    location: Optional[str] = Field(None, description="Current location of the candidate")
    current_company: Optional[str] = Field(None, description="Current company of the candidate")
    current_position: Optional[str] = Field(None, description="Current position of the candidate")
    current_remuneration: Optional[str] = Field(None, description="Current remuneration of the candidate")
    expected_remuneration: Optional[str] = Field(None, description="Expected remuneration of the candidate")
    notice_period: Optional[str] = Field(None, description="Notice period of the candidate")
    current_wfh_model: Optional[str] = Field(None, description="Current work from home model of the candidate")
    wfh_model_considered: Optional[str] = Field(None, description="Work from home model the candidate would consider")

class CareerOverviewItem(BaseModel):
    company_name: Optional[str] = Field(None, description="Name of the company")
    position: Optional[str] = Field(None, description="Position held at the company")
    start_date: Optional[str] = Field(None, description="Start date of employment")
    end_date: Optional[str] = Field(None, description="End date of employment")
    location: Optional[str] = Field(None, description="Location of the company")
    description: Optional[str] = Field(None, description="Description of the role and responsibilities")

class QualificationItem(BaseModel):
    degree: Optional[str] = Field(None, description="Degree obtained")
    institution: Optional[str] = Field(None, description="Educational institution")
    location: Optional[str] = Field(None, description="Location of the institution")
    year: Optional[str] = Field(None, description="Year of graduation")
    minor: Optional[str] = Field(None, description="Minor subject")
    modules: Optional[List[str]] = Field(None, description="List of relevant modules or coursework")

class CVSchema(BaseModel):
    personal_information: Optional[PersonalInformation] = Field(None, description="Personal information of the candidate")
    executive_brief: Optional[str] = Field(None, description="Executive summary or brief of the candidate")
    qualifications: Optional[List[QualificationItem]] = Field(None, description="List of qualifications of the candidate")
    career_overview: Optional[List[CareerOverviewItem]] = Field(None, description="List of career overview items")
    skills: Optional[List[str]] = Field(None, description="List of skills of the candidate")

# Configure Gemini API
GOOGLE_API_KEY = "AIzaSyDTQWHApZj47R-KWX64-FKtruz_dfK6WFw"
genai.configure(api_key=GOOGLE_API_KEY)
logging.info("Gemini API key configured directly in the script. This is not recommended for production.")

# Configure the global model to expect text responses
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash-exp",
    generation_config=genai.GenerationConfig(
        temperature=0.2,
        top_p=0.95,
        top_k=40,
        max_output_tokens=8192,
    ),
)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    logging.info(f"Extracting text from PDF: {pdf_path}")
    text = ""
    try:
        with open(pdf_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
        logging.info(f"Successfully extracted text from PDF. Length: {len(text)} characters.")
    except Exception as e:
        logging.error(f"Error extracting text from PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    logging.info(f"Extracting text from DOCX: {docx_path}")
    text = ""
    try:
        doc = Document(docx_path)
        text += "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += f"{cell.text.strip()}\n"
        logging.info(f"Successfully extracted text from DOCX. Length: {len(text)} characters.")
    except Exception as e:
        logging.error(f"Error extracting text from DOCX {docx_path}: {e}")
    return text

def extract_text_from_txt(txt_path):
    logging.info(f"Extracting text from TXT: {txt_path}")
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logging.error(f"Error reading TXT file {txt_path}: {e}")
        return ""

def preprocess_response_data(raw_data: dict) -> dict:
    """
    Preprocesses raw data to ensure it matches the schema.
    Converts comma-separated strings in 'modules' to a list of strings.
    """
    if 'qualifications' in raw_data:
        for qualification in raw_data['qualifications']:
            if 'modules' in qualification and isinstance(qualification['modules'], str):
                # Convert comma-separated string to list
                qualification['modules'] = [module.strip() for module in qualification['modules'].split(',')]
    return raw_data

def format_structured_data(structured_data: CVSchema) -> str:
    """Formats the structured data into a human-readable string."""
    formatted_text = "## Extracted CV Data:\n\n"

    if structured_data.personal_information:
        formatted_text += "### Personal Information:\n"
        for key, value in structured_data.personal_information.model_dump(exclude_none=True).items():
            formatted_text += f"- **{key.replace('_', ' ').title()}:** {value}\n"
        formatted_text += "\n"

    if structured_data.executive_brief:
        formatted_text += f"### Executive Brief:\n{structured_data.executive_brief}\n\n"

    if structured_data.qualifications:
        formatted_text += "### Qualifications:\n"
        for qualification in structured_data.qualifications:
            parts = [f"**{key.replace('_', ' ').title()}**: {value}" for key, value in qualification.model_dump(exclude_none=True).items()]
            formatted_text += f"- {', '.join(parts)}\n"
        formatted_text += "\n"

    if structured_data.career_overview:
        formatted_text += "### Career Overview:\n"
        for career in structured_data.career_overview:
            formatted_text += f"- **{career.position}** at **{career.company_name}** ({career.start_date} - {career.end_date}): {career.description}\n"
        formatted_text += "\n"

    if structured_data.skills:
        formatted_text += "### Skills:\n"
        for skill in structured_data.skills:
            formatted_text += f"- {skill}\n"
        formatted_text += "\n"

    return formatted_text

def format_ai_summary(ai_summary_text):
    formatted_summary = ""
    lines = ai_summary_text.split("___")
    for section in lines:
        section = section.strip()
        if not section:
            continue
        formatted_summary += f"{section}\n\n"
    return formatted_summary

def generate_career_brief(structured_data):
    if not structured_data:
        return None
    personal_info = structured_data.personal_information
    career_overview = structured_data.career_overview
    qualifications = structured_data.qualifications
    skills = structured_data.skills

    full_name = personal_info.full_name if personal_info and personal_info.full_name else "Unknown"

    # Calculate total years of experience (basic approximation)
    total_experience_years = 0
    if career_overview:
        most_recent_start_year = None
        earliest_end_year = None
        for job in career_overview:
            start_year_str = job.start_date.split()[-1] if job.start_date else None
            end_year_str = job.end_date.split()[-1] if job.end_date else None
            try:
                start_year = int(start_year_str) if start_year_str and start_year_str.isdigit() else None
                end_year = int(end_year_str) if end_year_str and (end_year_str.lower() == 'present' or end_year_str.isdigit()) else None

                if start_year:
                    if most_recent_start_year is None or start_year > most_recent_start_year:
                        most_recent_start_year = start_year
                if end_year:
                    if earliest_end_year is None or end_year < earliest_end_year:
                        earliest_end_year = end_year
            except ValueError:
                continue
        if most_recent_start_year and earliest_end_year:
            total_experience_years = most_recent_start_year - earliest_end_year

    qualifications_str = ", ".join([q.degree for q in qualifications if q.degree]) if qualifications else ""
    skills_str = ", ".join(skills) if skills else ""

    prompt = f"""
    Instructions:

    Provide a concise career brief for the candidate, {full_name}, who has approximately {total_experience_years} years of relevant industry experience.

    Start with a brief introduction summarizing their key qualifications ({qualifications_str}) and skills ({skills_str}).

    Then, list their recent work history from the past six years in reverse chronological order, including the position and company name. For each role, briefly describe their responsibilities and contributions.

    Following the work history, summarize any notable accomplishments.

    Finally, list their technical skills and methodologies.

    Ensure the career brief is informative and easy to read, avoiding jargon and focusing on clear, factual statements. Do not include specific dates beyond years and durations, names of educational institutions, or personal details unrelated to professional qualifications.
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        logging.error(f"Error generating career brief: {e}")
        st.error("Failed to generate career brief. Please check the logs for details.")
        return ""

# Streamlit app
def main():
    st.title("Candidate Brief Generator")

    st.warning("You have set the Gemini API key directly in the script. This is NOT recommended for production due to security risks. Consider using environment variables or Streamlit secrets instead.")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.subheader("Upload CV")
        uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, TXT):", type=["pdf", "docx", "txt"])

    with col2:
        st.subheader("Structured CV Data (JSON)")
        with st.expander("Show/Hide JSON Data", expanded=False):
            if 'structured_data_json' in st.session_state:
                st.json(st.session_state['structured_data_json'])
                st.download_button(
                    label="Download Structured Data (JSON)",
                    data=st.session_state['structured_data_json'],
                    file_name="candidate_data.json",
                    mime="application/json"
                )

    with col3:
        st.subheader("Structured CV Data (Human Readable)")
        with st.expander("Show/Hide Readable Data", expanded=False):
            if 'formatted_data' in st.session_state:
                st.markdown(st.session_state['formatted_data'])
                st.download_button(
                    label="Download Structured Data (TXT)",
                    data=st.session_state['formatted_data'],
                    file_name="candidate_data_formatted.txt",
                    mime="text/plain"
                )

    st.markdown("---")  # Separator for the new section
    st.subheader("AI Generated Career Brief Summary")

    if st.button("Generate Career Brief"):
         if 'structured_data' in st.session_state and st.session_state['structured_data']:
            ai_summary_text = generate_career_brief(st.session_state['structured_data'])
            st.session_state['ai_summary'] = ai_summary_text
         else:
            st.warning("Please upload and process a CV first to generate the career brief.")

    if 'ai_summary' in st.session_state:
        with st.expander("Career Brief Summary", expanded=False):
            st.markdown(st.session_state['ai_summary'])

    if uploaded_file is not None:
        try:
            temp_file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
            with open(temp_file_path, 'wb') as temp_file:
                temp_file.write(uploaded_file.read())

            file_extension = uploaded_file.name.rsplit('.', 1)[1].lower()
            if file_extension == 'pdf':
                extracted_text = extract_text_from_pdf(temp_file_path)
            elif file_extension == 'docx':
                extracted_text = extract_text_from_docx(temp_file_path)
            elif file_extension == 'txt':
                extracted_text = extract_text_from_txt(temp_file_path)
            else:
                st.error("Unsupported file type.")
                return

            if not extracted_text.strip():
                st.error("The file does not contain valid readable content.")
                return

            prompt = f"""
            Instructions:
            Analyze the following CV text and extract the key information into a JSON object based on the provided schema. Ensure all fields in the schema are present in the JSON. For the 'modules' field within 'qualifications', please provide a list of strings. If a piece of information is not found, use an empty string or an empty list as appropriate.

            Here is the schema:
            {CVSchema.model_json_schema()}

            Here is the CV text:
            {extracted_text}
            """

            generation_config = genai.GenerationConfig(
                response_mime_type="application/json",
                # response_schema=CVSchema # Removing response_schema to handle potential string output for modules
            )

            with st.spinner("Processing CV and extracting structured data..."):
                response = model.generate_content(prompt, generation_config=generation_config)
                full_response = response.text
                logging.info(f"Full response from Gemini API: {full_response}")

            try:
                raw_response_data = json.loads(full_response)
                processed_response_data = preprocess_response_data(raw_response_data)

                structured_data = CVSchema.model_validate(processed_response_data)
                st.session_state['structured_data'] = structured_data
                st.session_state['structured_data_json'] = json.dumps(structured_data.model_dump(exclude_none=True), indent=4)
                logging.info(f"Successfully parsed JSON: {structured_data}")

                formatted_data = format_structured_data(structured_data)
                st.session_state['formatted_data'] = formatted_data
            except json.JSONDecodeError as e:
                st.error("Failed to parse JSON from Gemini API response.")
                logging.error(f"JSON parsing error: {e}")
                st.error(f"Raw response from API: '{full_response}'")
            except Exception as e:
                st.error(f"An error occurred while parsing the response: {e}")
                logging.error(f"Error parsing Gemini API response: {e}")

        except Exception as e:
            logging.error(f"Error processing file: {e}")
            st.error(f"An error occurred while processing the file: {e}")

if __name__ == "__main__":
    main()